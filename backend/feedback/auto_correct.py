from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from bs4 import BeautifulSoup
from io import BytesIO

# --- FUNCTION 1: Keep this so main.py doesn't crash ---
def improve_docx(file_bytes, suggestions):
    """
    Applies automated fixes to the original document bytes.
    """
    doc = Document(BytesIO(file_bytes))
    # ... your existing logic for automated fixes ...
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

# --- FUNCTION 2: High-Fidelity Download Logic ---
def html_to_docx(html_content):
    """
    Converts edited HTML back to a professional DOCX with formatting.
    """
    doc = Document()
    
    # Standard professional margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    soup = BeautifulSoup(html_content, "html.parser")
    
    for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'div', 'hr']):
        # Handle Horizontal Lines (simulating your original resume layout)
        if element.name == 'hr':
            p = doc.add_paragraph()
            # Logic to add a paragraph border as a line
            continue

        text = element.get_text().strip()
        if not text:
            continue

        p = doc.add_paragraph()
        if element.name == 'h1': # Name/Header
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(18)
        else:
            p.add_run(text)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()